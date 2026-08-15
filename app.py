import os
import re
import csv
import time
import shutil
import base64
import logging
import zipfile
import urllib.parse
from io import BytesIO
from pathlib import Path

import requests
import openpyxl
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image as PILImage
import streamlit as st

CONFIG = {
    'output_dir': 'pubchem_output',
    'image_folder': 'images',
    'sdf_folder': 'sdf_3d',
    'excel_out_name': 'ARoHaN_Lab_CompoundData.xlsx',
    'docx_out_name': 'ARoHaN_Lab_CompoundData.docx',
    'csv_out_name': 'ARoHaN_Lab_CompoundData.csv',
    'zip_out_name': 'ARoHaN_Lab_CompoundData_results.zip',
    'compound_column': 'A',
    'header_row': 1,
    'request_delay': 0.25,
    'request_timeout': 25,
    'max_retries': 3,
    'max_compounds': 1500
}

PUG_URL = 'https://pubchem.ncbi.nlm.nih.gov/rest/pug'
CAS_RE = re.compile(r'^\d{2,7}-\d{2}-\d$')
INCHIKEY_RE = re.compile(r'^[A-Z]{14}-[A-Z]{10}-[A-Z]$')

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger('pubchem_extractor')

def setup_workspace():
    base_dir = Path(CONFIG['output_dir'])
    if base_dir.exists():
        shutil.rmtree(base_dir)
    paths = {
        'base': base_dir,
        'images': base_dir / CONFIG['image_folder'],
        'sdf': base_dir / CONFIG['sdf_folder'],
    }
    for p in paths.values():
        p.mkdir(parents=True, exist_ok=True)
    return paths

def sanitize_filename(name):
    return re.sub(r'[\\/*?:"<>| ]', '_', str(name))[:50]

def detect_identifier_type(term):
    term = str(term).strip()
    if CAS_RE.match(term):
        return 'name' 
    if INCHIKEY_RE.match(term):
        return 'inchikey'
    return 'name'

def query_pubchem_api(url, retry_count=0):
    try:
        time.sleep(CONFIG['request_delay'])
        response = requests.get(url, timeout=CONFIG['request_timeout'])
        if response.status_code == 200:
            return response
        elif response.status_code in [400, 404]:
            return None
    except Exception as e:
        if retry_count < CONFIG['max_retries']:
            time.sleep(2 ** retry_count)
            return query_pubchem_api(url, retry_count + 1)
    return None

def resolve_compound_cid(identifier):
    id_type = detect_identifier_type(identifier)
    encoded_id = urllib.parse.quote(str(identifier).strip())
    url = f'{PUG_URL}/compound/{id_type}/{encoded_id}/cids/JSON'
    res = query_pubchem_api(url)
    if res:
        try:
            data = res.json()
            return data['IdentifierList']['CID'][0]
        except (KeyError, IndexError):
            return None
    return None

def fetch_compound_properties(cid):
    props = 'MolecularFormula,MolecularWeight,SMILES,IUPACName,XLogP,TPSA,HBondDonorCount,HBondAcceptorCount'
    url = f'{PUG_URL}/compound/cid/{cid}/property/{props}/JSON'
    res = query_pubchem_api(url)
    if res:
        try:
            return res.json()['PropertyTable']['Properties'][0]
        except (KeyError, IndexError):
            return {}
    return {}

def download_3d_sdf(cid, dest_path):
    url = f'{PUG_URL}/compound/cid/{cid}/SDF?record_type=3d'
    res = query_pubchem_api(url)
    if res and len(res.content) > 100:
        with open(dest_path, 'wb') as f:
            f.write(res.content)
        return True
    return False

def download_2d_image(cid, dest_path):
    url = f'{PUG_URL}/compound/cid/{cid}/PNG'
    res = query_pubchem_api(url)
    if res:
        with open(dest_path, 'wb') as f:
            f.write(res.content)
        return True
    return False

def generate_excel_report(rows, paths):
    wb = openpyxl.Workbook()
    ws_summary = wb.active
    ws_summary.title = 'Summary'
    ws_summary.append(['ARoHaN Lab - Batch Processing Log'])
    ws_summary.append(['Timestamp', time.strftime('%Y-%m-%d %H:%M:%S')])
    ws_summary.append(['Total Checked', len(rows)])
    ws_summary.append(['Successfully Extracted', len([r for r in rows if r['status'] == 'Success'])])
    ws_summary.append(['Failed/Not Found', len([r for r in rows if r['status'] != 'Success'])])
    ws_data = wb.create_sheet(title='Molecular Properties')
    headers = ['Input Query', 'Status', 'CID', 'IUPAC Name', 'Formula', 'Molecular Weight', 'LogP', 'TPSA', 'H-Bond Donors', 'H-Bond Acceptors', 'SMILES', '2D Structure']
    ws_data.append(headers)
    for idx, r in enumerate(rows, start=2):
        ws_data.append([r['query'], r['status'], r['cid'], r['name'], r['formula'], r['mw'], r['logp'], r['tpsa'], r['donors'], r['acceptors'], r['smiles'], ''])
        if r['status'] == 'Success' and r['img_path'] and os.path.exists(r['img_path']):
            try:
                img = XLImage(r['img_path'])
                img.width = 110
                img.height = 110
                ws_data.row_dimensions[idx].height = 90
                ws_data.add_image(img, f'L{idx}')
            except Exception:
                pass
    ws_data.column_dimensions['L'].width = 18
    wb.save(paths['base'] / CONFIG['excel_out_name'])

def generate_word_report(rows, paths):
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(14.0)
    section.page_height = Inches(8.5)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    title = doc.add_paragraph()
    run = title.add_run('ARoHaN Lab — Molecular Data Extraction Matrix')
    run.font.size = Pt(18)
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Generated on: {time.strftime("%Y-%m-%d %H:%M:%S")} | PI: Dr. Manik Ghosh · BIT Mesra\n')
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    tblLayout = parse_xml(r'<w:tblLayout %s w:type="fixed"/>' % nsdecls('w'))
    tblPr.append(tblLayout)
    col_widths = [Cm(4.0), Cm(2.25), Cm(9.5), Cm(3.0), Cm(2.0), Cm(9.0), Cm(3.45)]
    hdr_cells = table.rows[0].cells
    hdrs = ['Query', 'CID', 'Chemical IUPAC Name', 'Formula', 'Mol. Wt', 'SMILES String', '2D Structure']
    for i, title_text in enumerate(hdrs):
        hdr_cells[i].text = title_text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(r'<w:tblHeader %s/>' % nsdecls('w')))
    for r in rows:
        if r['status'] == 'Success':
            row_cells = table.add_row().cells
            row_cells[0].text = str(r['query'])
            row_cells[1].text = str(r['cid'])
            row_cells[2].text = str(r['name'])
            row_cells[3].text = str(r['formula'])
            row_cells[4].text = str(r['mw'])
            row_cells[5].text = str(r['smiles'])
            if r['img_path'] and os.path.exists(r['img_path']):
                p = row_cells[6].paragraphs[0]
                p.add_run().add_picture(str(r['img_path']), width=Inches(1.2))
    for row in table.rows:
        r_trPr = row._tr.get_or_add_trPr()
        r_trPr.append(parse_xml(r'<w:cantSplit %s/>' % nsdecls('w')))
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
    doc.save(paths['base'] / CONFIG['docx_out_name'])

def generate_csv_report(rows, paths):
    csv_path = paths['base'] / CONFIG['csv_out_name']
    fields = ['query', 'status', 'cid', 'name', 'formula', 'mw', 'logp', 'tpsa', 'donors', 'acceptors', 'smiles']
    with open(csv_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction='ignore')
        writer.writeheader()
        for r in rows:
            writer.writerow(r)

def run_batch_extraction_pipeline(input_excel_path):
    paths = setup_workspace()
    wb = openpyxl.load_workbook(input_excel_path, data_only=True)
    sheet = wb.active
    queries = []
    for row in range(CONFIG['header_row'] + 1, sheet.max_row + 1):
        val = sheet[f"{CONFIG['compound_column']}{row}"].value
        if val and str(val).strip():
            queries.append(str(val).strip())
    unique_queries = list(dict.fromkeys(queries))[:CONFIG['max_compounds']]
    st_progress = st.progress(0)
    st_status = st.empty()
    compiled_results = []
    total = len(unique_queries)
    for idx, query in enumerate(unique_queries):
        st_status.text(f'Processing ({idx+1}/{total}): {query}')
        st_progress.progress(int(((idx) / total) * 100))
        result = {
            'query': query, 'status': 'Not Found', 'cid': '', 'name': '', 'formula': '',
            'mw': '', 'logp': '', 'tpsa': '', 'donors': '', 'acceptors': '', 'smiles': '', 'img_path': ''
        }
        cid = resolve_compound_cid(query)
        if cid:
            result['cid'] = cid
            props = fetch_compound_properties(cid)
            if props:
                result['status'] = 'Success'
                result['name'] = props.get('IUPACName', 'N/A')
                result['formula'] = props.get('MolecularFormula', 'N/A')
                result['mw'] = props.get('MolecularWeight', 'N/A')
                result['logp'] = props.get('XLogP', 'N/A')
                result['tpsa'] = props.get('TPSA', 'N/A')
                result['donors'] = props.get('HBondDonorCount', 'N/A')
                result['acceptors'] = props.get('HBondAcceptorCount', 'N/A')
                result['smiles'] = props.get('SMILES', 'N/A')
                safe_name = sanitize_filename(query)
                img_p = paths['images'] / f'{safe_name}_{cid}.png'
                parts_sdf = paths['sdf'] / f'{safe_name}_{cid}.sdf'
                if download_2d_image(cid, img_p):
                    result['img_path'] = str(img_p)
                download_3d_sdf(cid, parts_sdf)
        compiled_results.append(result)
    st_progress.progress(100)
    st_status.text('Building final report packages...')
    final_ordered_rows = []
    for orig_q in queries:
        match = next((item for item in compiled_results if item['query'] == orig_q), None)
        if match:
            final_ordered_rows.append(match)
    generate_excel_report(final_ordered_rows, paths)
    generate_word_report(final_ordered_rows, paths)
    generate_csv_report(final_ordered_rows, paths)
    zip_path = paths['base'].parent / CONFIG['zip_out_name']
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, _, files in os.walk(paths['base']):
            for file in files:
                file_p = Path(root) / file
                zipf.write(file_p, file_p.relative_to(paths['base']))
    return zip_path

def main():
    st.set_page_config(page_title='ARoHaN Lab - Data Extractor', page_icon='🧬')
    st.title('🧬 ARoHaN Lab')
    st.subheader('Advance Research on Herbals and Naturals — Compound Data Extractor')
    st.markdown('**Principal Investigator:** Dr. Manik Ghosh · BIT Mesra')
    st.divider()
    st.markdown('''
    ### 📊 Instructions:
    1. Upload a standard Excel spreadsheet (`.xlsx`).
    2. Ensure your target chemical compounds or CAS strings are placed in **Column A**.
    3. Ensure your row list begins immediately on **Row 2** (Row 1 is assumed to be your header).
    ''')
    uploaded_file = st.file_uploader('Choose your Excel configuration spreadsheet', type=['xlsx'])
    if uploaded_file is not None:
        st.success('Spreadsheet read safely into memory cache.')
        if st.button('🚀 Run Batch Extraction Pipeline'):
            temp_input_path = 'backend_temp_input.xlsx'
            with open(temp_input_path, 'wb') as f:
                f.write(uploaded_file.getbuffer())
            with st.spinner('Executing secure PubChem crawlers... Please leave this tab open.'):
                try:
                    archive_out = run_batch_extraction_pipeline(temp_input_path)
                    with open(archive_out, 'rb') as f:
                        st.balloons()
                        st.success('🎉 Processing complete! Download link compiled below.')
                        st.download_button(
                            label='💾 Download Extract Archive (ZIP)',
                            data=f,
                            file_name=CONFIG['zip_out_name'],
                            mime='application/zip'
                        )
                except Exception as e:
                    st.error(f'A processing execution anomaly occurred: {e}')
                finally:
                    if os.path.exists(temp_input_path):
                        os.remove(temp_input_path)

if __name__ == '__main__':
    main()